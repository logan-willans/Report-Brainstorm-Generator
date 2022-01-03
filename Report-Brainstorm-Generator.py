import docx, datetime, os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
client_reference = input("Enter client reference: ")
doc = docx.Document()
categories = ["Alarm System", "Security Cameras", "Perimeter Detection", "Exit\n"]
alarmSystemCategories = ["No Existing Alarm System", "Existing Elk System", "Existing ADT System", "Existing Vivint System", "Return to Main Menu\n"]
securityCameraCategories = ["Thermal Cameras", "Pan-Tilt-Zoom Cameras", "Consumer-Grade Cameras", "Multi-Sensor Cameras", "Return to Main Menu\n"]
perimeterDetectionCategories = ["Infrared Beams", "LIDAR", "Radar", "Motion Viewers", "Return to Main Menu\n"]

selection = 0
choice = 0

def prints_main_menu():
    print("\nThis program can pull text from the categories listed below. Enter '4' to exit the program.\n")
    for i, x in enumerate(categories, 1):
        print(i, "-", x)
    while True:
        try:
            global selection
            selection = int(input("Please enter the integer value of the option you're interested in: "))
            while selection > 13 or selection < 1:
                selection = int(input("Input was outside of specified range. Please try again: "))
            break
        except ValueError:
            print("Only integer values are accepted. Please try again.")

def alarmSystemMenu():
    global choice
    print("\nHere are the text options for alarm systems.\n")
    for i, x in enumerate(alarmSystemCategories, 1):
        print(i, "-", x)
    while True:
        try:
            choice = int(input("Please enter the integer value of the option you're interested in: "))
            while choice > 5 or choice < 1:
                choice = int(input("Input was outside of specified range. Please try again: "))
            break
        except ValueError:
            print("Only integer values are accepted. Please try again.")

def securityCameraMenu():
    global choice
    print("\nHere are the text options for security cameras.\n")
    for i, x in enumerate(securityCameraCategories, 1):
        print(i, "-", x)
    while True:
        try:
            choice = int(input("Please enter the integer value of the option you're interested in: "))
            while choice > 5 or choice < 1:
                choice = int(input("Input was outside of specified range. Please try again: "))
            break
        except ValueError:
            print("Only integer values are accepted. Please try again.")

def perimeterDetectionMenu():
    global choice
    print("\nHere are the text options for perimeter detection.\n")
    for i, x in enumerate(perimeterDetectionCategories, 1):
        print(i, "-", x)
    while True:
        try:
            choice = int(input("Please enter the integer value of the option you're interested in: "))
            while choice > 5 or choice < 1:
                choice = int(input("Input was outside of specified range. Please try again: "))
            break
        except ValueError:
            print("Only integer values are accepted. Please try again.")

def title_page():
    today = str(datetime.date.today())
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    first_header = doc.add_heading(client_reference + " Brainstorm", 1)
    first_header_format = first_header.paragraph_format
    first_header_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    third_header = doc.add_paragraph(today)
    third_header_format = third_header.paragraph_format
    third_header_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    third_header.bold=True
    doc.add_page_break()

title_page()
prints_main_menu()

while selection != 4:
    if selection == 1:
        alarmSystemMenu()
        while choice != 5:
            if choice == 1:
                print("\nText for option 1 has been inserted.")
                doc.add_paragraph().add_run("No Existing Alarm System - Reference Example Report Page XX").bold=True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                alarmSystemMenu()
            if choice == 2:
                print("\nText for option 2 has been inserted.")
                doc.add_paragraph().add_run("Existing Elk System - Reference Example Report Page XX").bold = True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                alarmSystemMenu()
            if choice == 3:
                print("\nText for option 3 has been inserted.")
                doc.add_paragraph().add_run("Existing ADT System - Reference Example Report Page XX").bold = True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                alarmSystemMenu()
            if choice == 4:
                print("\nText for option 4 has been inserted.")
                doc.add_paragraph().add_run("Existing Vivent System - Reference Example Report Page XX").bold = True
                doc.add_paragraph(
                    "Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                alarmSystemMenu()
        prints_main_menu()

    if selection == 2:
        securityCameraMenu()
        while choice != 5:
            if choice == 1:
                print("\nText for option 1 has been inserted.")
                doc.add_paragraph().add_run("Thermal Cameras - Reference Example Report Page XX").bold=True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                securityCameraMenu()
            if choice == 2:
                print("\nText for option 2 has been inserted.")
                doc.add_paragraph().add_run("Pan-Tilt-Zoom Cameras - Reference Example Report Page XX").bold = True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                securityCameraMenu()
            if choice == 3:
                print("\nText for option 3 has been inserted.")
                doc.add_paragraph().add_run("Consumer-Grade Cameras - Reference Example Report Page XX").bold = True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                securityCameraMenu()
            if choice == 4:
                print("\nText for option 4 has been inserted.")
                doc.add_paragraph().add_run("Multi-Sensor Cameras - Reference Example Report Page XX").bold = True
                doc.add_paragraph(
                    "Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                securityCameraMenu()
        prints_main_menu()

    #perimeterDetectionCategories = ["Infrared Beams", "LIDAR", "Radar", "Motion Viewers", "Return to Main Menu\n"]
    if selection == 3:
        perimeterDetectionMenu()
        while choice != 5:
            if choice == 1:
                print("\nText for option 1 has been inserted.")
                doc.add_paragraph().add_run("Infrared Beams - Reference Example Report Page XX").bold=True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                perimeterDetectionMenu()
            if choice == 2:
                print("\nText for option 2 has been inserted.")
                doc.add_paragraph().add_run("LIDAR - Reference Example Report Page XX").bold = True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                perimeterDetectionMenu()
            if choice == 3:
                print("\nText for option 3 has been inserted.")
                doc.add_paragraph().add_run("Radar - Reference Example Report Page XX").bold = True
                doc.add_paragraph("Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                perimeterDetectionMenu()
            if choice == 4:
                print("\nText for option 4 has been inserted.")
                doc.add_paragraph().add_run("Motion Viewers - Reference Example Report Page XX").bold = True
                doc.add_paragraph(
                    "Assum numquam te eum, cu option delicatissimi pri. Officiis concludaturque duo id. Sit rationibus appellantur in, ad agam prompta laoreet per. Quo tollit scaevola accusamus cu, diam diceret ei eos, nullam maiorum ne usu. Has posse corpora et, putent iisque eam at, alia paulo philosophia ad sit.")
                perimeterDetectionMenu()
        prints_main_menu()

doc.save(client_reference.replace(" ", "_") + "_Residential" + "_Assessment" + ".docx")
os.system(client_reference.replace(" ", "_") + "_Residential" + "_Assessment" + ".docx")